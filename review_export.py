"""Review export: render a survey schema as a .docx for offline review.

Produces a Word document showing everything a reviewer needs and nothing a
participant sees differently: verbatim texts in the primary language, answer
formats, page/question visibility logic in plain English, randomization pools
(from the database, not the schema), and per-element flags for which
translations exist.

Pure module: no FastAPI, no database. The route hands in the parsed schema and
the pool rows; this builds bytes.
"""

import json
import re
from datetime import date

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Twips

# Mirror of the locale set in templates/survey.html — keep the two in sync.
LANGS = ["en", "de", "fr", "it"]

GRAY = RGBColor(0x66, 0x66, 0x66)
NAVY = RGBColor(0x2B, 0x4A, 0x73)
RED = RGBColor(0xB0, 0x30, 0x30)
BOX_FILL = "F4F4F4"
BOX_BORDER = "999999"


# --- localization helpers ---

def _is_loc_obj(v) -> bool:
    """A SurveyJS localizable string rendered as {locale: text, ...}."""
    if not isinstance(v, dict) or not v:
        return False
    return all(isinstance(x, str) for x in v.values()) and any(
        k == "default" or k in LANGS for k in v
    )


def loc_text(v) -> str:
    """Primary-language text of a plain or localized string."""
    if isinstance(v, str):
        return v
    if _is_loc_obj(v):
        return v.get("default") or v.get("en") or next(iter(v.values()))
    return "" if v is None else str(v)


def _locales_in(obj) -> set:
    """Every locale key used anywhere inside `obj` ('default' counts as 'en')."""
    found = set()

    def walk(v):
        if isinstance(v, list):
            for x in v:
                walk(x)
        elif isinstance(v, dict):
            if _is_loc_obj(v):
                for k in v:
                    if k in LANGS:
                        found.add(k)
                    elif k == "default":
                        found.add("en")
            else:
                for x in v.values():
                    walk(x)

    walk(obj)
    return found


# --- expression humanizer ---

def humanize_expr(expr: str) -> str:
    """Best-effort plain-English rendering of a SurveyJS visibleIf expression."""
    s = expr
    s = re.sub(r"\{([^}]+)\}", r"\1", s)
    for a, b in [
        (" notcontains ", " does not include "),
        (" contains ", " includes "),
        (" anyof ", " is any of "),
        (" notempty", " has an answer"),
        (" empty", " is unanswered"),
        (" <> ", " is not "),
        (" >= ", " is at least "),
        (" <= ", " is at most "),
        (" = ", " is "),
    ]:
        s = s.replace(a, b)
    return s


# --- low-level docx helpers ---

def _shade(paragraph, fill=BOX_FILL, border=None):
    """Shaded quote-box look: background fill and optional left border."""
    p_pr = paragraph._p.get_or_add_pPr()
    if border:
        pbdr = OxmlElement("w:pBdr")
        left = OxmlElement("w:left")
        left.set(qn("w:val"), "single")
        left.set(qn("w:sz"), "24")
        left.set(qn("w:space"), "8")
        left.set(qn("w:color"), border)
        pbdr.append(left)
        p_pr.append(pbdr)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def _runs_from_html_inline(paragraph, html, bold=False):
    """Emit runs for inline HTML: <strong>/<em>/<b>/<i>, entities, tags stripped."""
    ital = False
    for part in re.split(r"(</?strong>|</?em>|</?b>|</?i>)", html):
        if part in ("<strong>", "<b>"):
            bold = True
        elif part in ("</strong>", "</b>"):
            bold = False
        elif part in ("<em>", "<i>"):
            ital = True
        elif part in ("</em>", "</i>"):
            ital = False
        elif part:
            text = re.sub(r"<[^>]+>", "", part)
            text = (
                text.replace("&nbsp;", " ")
                .replace("&amp;", "&")
                .replace("&lt;", "<")
                .replace("&gt;", ">")
            )
            if text:
                run = paragraph.add_run(text)
                run.bold = bold
                run.italic = ital


def _note(doc, text):
    """Gray italic logic note — content participants never see."""
    p = doc.add_paragraph()
    run = p.add_run("▸ " + text)
    run.italic = True
    run.font.color.rgb = GRAY
    run.font.size = Pt(9)
    return p


def _answer_line(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Twips(360)
    run = p.add_run(text)
    run.font.color.rgb = GRAY
    run.font.size = Pt(9)
    return p


def _choice_lines(doc, choices, symbol):
    for c in choices:
        text = loc_text(c.get("text", c.get("value")) if isinstance(c, dict) else c)
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Twips(360)
        p.paragraph_format.space_after = Pt(2)
        p.add_run(f"{symbol}  {text}")


def _render_html_element(doc, html):
    """Block-level mini renderer: <p> and <li> inside an optional boxed <div>."""
    boxed = "<div" in html
    blocks = re.findall(r"<p([^>]*)>(.*?)</p>|<li[^>]*>(.*?)</li>", html, re.S)
    for attrs, p_html, li_html in blocks:
        is_li = bool(li_html)
        content = li_html if is_li else p_html
        para = doc.add_paragraph()
        para.paragraph_format.space_after = Pt(4)
        if boxed:
            _shade(para, border=BOX_BORDER)
            para.paragraph_format.left_indent = Twips(480 if is_li else 240)
            para.paragraph_format.right_indent = Twips(240)
        if is_li:
            para.add_run("•  ")
        header = "font-weight:600" in attrs
        _runs_from_html_inline(para, content, bold=header)
    if boxed:
        doc.add_paragraph()


# --- per-question rendering ---

def _rating_scale_text(el) -> str:
    if el.get("rateValues"):
        values = " – ".join(loc_text(v.get("text", v.get("value")) if isinstance(v, dict) else v)
                            for v in el["rateValues"])
    else:
        lo = el.get("rateMin", 1)
        hi = el.get("rateMax", 5)
        values = " – ".join(str(n) for n in range(lo, hi + 1))
    text = f"Answer: rating scale {values}"
    lo_desc = loc_text(el.get("minRateDescription", ""))
    hi_desc = loc_text(el.get("maxRateDescription", ""))
    if lo_desc or hi_desc:
        text += f"   (lowest = {lo_desc}; highest = {hi_desc})"
    return text


def _matrix_table(doc, el):
    columns = [
        loc_text(c.get("text", c.get("value")) if isinstance(c, dict) else c).replace("\n", " — ")
        for c in el.get("columns", [])
    ]
    table = doc.add_table(rows=1, cols=len(columns) + 1)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, text in enumerate([""] + columns):
        p = hdr[i].paragraphs[0]
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(8)
        if i:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for row in el.get("rows", []):
        cells = table.add_row().cells
        run = cells[0].paragraphs[0].add_run(
            loc_text(row.get("text", row.get("value")) if isinstance(row, dict) else row)
        )
        run.font.size = Pt(9)
        for cell in cells[1:]:
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run("○").font.size = Pt(9)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    doc.add_paragraph()


def _locale_counts(el):
    """(counts, total): how many of the element's localized strings carry each
    locale, and how many localized strings there are in total."""
    counts = {}
    total = 0

    def walk(v):
        nonlocal total
        if isinstance(v, list):
            for x in v:
                walk(x)
        elif isinstance(v, dict):
            if _is_loc_obj(v):
                total += 1
                for k in v:
                    loc = "en" if k == "default" else k
                    if loc in LANGS:
                        counts[loc] = counts.get(loc, 0) + 1
            else:
                for x in v.values():
                    walk(x)

    walk(el)
    return counts, total


def _translation_flag(el, survey_locales) -> str | None:
    """Per-element translation status vs. the locales used anywhere in the schema."""
    if len(survey_locales) <= 1:
        return None
    counts, total = _locale_counts(el)
    if total == 0:  # plain strings only: primary language, nothing translated
        others = ", ".join(l.upper() for l in sorted(survey_locales - {"en"}))
        return f"Translations: none (primary language only) — MISSING: {others}"
    parts = []
    for loc in sorted(counts):
        partial = " (partial)" if counts[loc] < total else ""
        parts.append(loc.upper() + partial)
    label = "Translations: " + ", ".join(parts)
    missing = survey_locales - set(counts)
    if missing:
        label += " — MISSING: " + ", ".join(l.upper() for l in sorted(missing))
    return label


def _render_question(doc, el, survey_locales):
    title = loc_text(el.get("title", el.get("name", "")))
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(title)
    run.bold = True
    if el.get("isRequired"):
        star = p.add_run(" *")
        star.bold = True
        star.font.color.rgb = RED

    if el.get("description"):
        d = doc.add_paragraph()
        run = d.add_run(loc_text(el["description"]))
        run.italic = True
        run.font.size = Pt(9)

    if el.get("visibleIf"):
        _note(doc, "Shown only if: " + humanize_expr(el["visibleIf"]))
    if el.get("requiredIf"):
        _note(doc, "Required only if: " + humanize_expr(el["requiredIf"]))
    if el.get("enableIf"):
        _note(doc, "Enabled only if: " + humanize_expr(el["enableIf"]))
    flag = _translation_flag(el, survey_locales)
    if flag:
        _note(doc, flag)

    kind = el.get("type")
    if kind == "rating":
        _answer_line(doc, _rating_scale_text(el))
    elif kind == "radiogroup":
        _choice_lines(doc, el.get("choices", []), "○")
    elif kind == "checkbox":
        _choice_lines(doc, el.get("choices", []), "☐")
        _answer_line(doc, "(multiple answers possible)")
    elif kind in ("dropdown", "tagbox"):
        if el.get("choicesByUrl"):
            _answer_line(doc, f"Answer: drop-down list loaded from {el['choicesByUrl'].get('url', '?')}")
        else:
            _answer_line(doc, "Answer: drop-down menu with the following options:")
            _choice_lines(doc, el.get("choices", []), "▾")
    elif kind == "boolean":
        yes = loc_text(el.get("labelTrue", "Yes"))
        no = loc_text(el.get("labelFalse", "No"))
        _answer_line(doc, f"Answer: {yes} / {no} switch")
    elif kind == "comment":
        _answer_line(doc, "Answer: open text box (free response).")
    elif kind == "text":
        hints = {"number": "numeric entry field", "date": "date picker", "email": "e-mail field"}
        _answer_line(doc, "Answer: " + hints.get(el.get("inputType"), "short text field") + ".")
    elif kind == "matrix":
        _matrix_table(doc, el)
    elif kind == "expression":
        _note(doc, f"Computed value (not asked): {el.get('expression', '')}")
    else:
        _note(doc, f"Question type '{kind}' not rendered here — check the JSON schema.")


def _render_element(doc, el, survey_locales):
    kind = el.get("type")
    if kind == "html":
        flag = _translation_flag(el, survey_locales)
        if flag:
            _note(doc, flag)
        _render_html_element(doc, loc_text(el.get("html", "")))
    elif kind in ("panel", "paneldynamic"):
        if el.get("title"):
            p = doc.add_paragraph()
            run = p.add_run(loc_text(el["title"]))
            run.bold = True
        if el.get("visibleIf"):
            _note(doc, "Panel shown only if: " + humanize_expr(el["visibleIf"]))
        for sub in el.get("elements", el.get("templateElements", [])):
            _render_element(doc, sub, survey_locales)
    else:
        _render_question(doc, el, survey_locales)


# --- randomization notes ---

def _pool_for_page(pools, page_name):
    for pool in pools:
        if page_name in pool["pool_pages"]:
            return pool
    return None


def _pool_page_note(pool, page_name) -> str:
    n = len(pool["pool_pages"])
    sc = pool["show_count"]
    text = (
        f"Randomized (pool “{pool['pool_name']}”): each participant sees {sc} of the "
        f"{n} pages in this pool, assigned so that all variants stay balanced."
    )
    if sc == 1 and pool.get("condition_var"):
        cmap = pool.get("condition_map") or {}
        value = cmap.get(page_name, page_name)
        text += f" When this page is drawn, the variable “{pool['condition_var']}” is set to “{value}”."
    elif sc == 1 and re.match(r"^info_([a-z])$", page_name, re.I):
        letter = page_name.split("_")[1].upper()
        text += f" When this page is drawn, the variable “condition” is set to “{letter}” (legacy naming convention)."
    return text


# --- document assembly ---

def build_review_docx(schema: dict, pools: list, survey_title: str = "") -> bytes:
    """Render `schema` (parsed SurveyJS JSON) + `pools` (rand_pools rows as dicts
    with parsed pool_pages/condition_map) into .docx bytes."""
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    survey_locales = _locales_in(schema) or {"en"}
    pages = schema.get("pages", [])
    n_questions = sum(
        1 for p in pages for e in p.get("elements", []) if e.get("type") != "html"
    )

    title = loc_text(schema.get("title")) or survey_title
    h = doc.add_heading(title, level=0)
    for run in h.runs:
        run.font.color.rgb = NAVY
    if schema.get("description"):
        p = doc.add_paragraph()
        run = p.add_run(loc_text(schema["description"]))
        run.italic = True
        run.font.size = Pt(13)

    meta = doc.add_paragraph()
    run = meta.add_run(
        f"Review export generated on {date.today().isoformat()} — "
        f"{len(pages)} pages, {n_questions} questions, "
        f"languages in schema: {', '.join(l.upper() for l in sorted(survey_locales))}."
    )
    run.font.color.rgb = GRAY
    run.font.size = Pt(9)

    doc.add_heading("How to read this document", level=1)
    guide = [
        ("Purpose. ", "This document reproduces the full content of the online "
         "questionnaire for offline review. The online layout differs (one page at "
         "a time, with a progress bar), but every text, question, and answer option "
         "shown here appears verbatim in the online version."),
        ("Shaded boxes ", "contain text participants read exactly as printed."),
        ("Notes in gray italics starting with ▸ ", "describe questionnaire logic "
         "(randomization, when a question appears, translation status). "
         "Participants never see them."),
        ("Answer formats ", "are described in small gray text under each question: "
         "○ marks single-choice options; ☐ marks multiple-choice options."),
        ("Asterisks (*) ", "mark questions that must be answered to continue."),
    ]
    if len(survey_locales) > 1:
        guide.append((
            "Languages. ",
            "Only the primary language is shown. Each element carries a note "
            "listing which translations exist and which are missing."
        ))
    for lead, rest in guide:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(lead).bold = True
        p.add_run(rest)

    if pools:
        doc.add_heading("Randomization", level=1)
        doc.add_paragraph(
            "Participants are randomly assigned pages from the pools below. "
            "Assignment is balanced: the least-used variant is always preferred. "
            "This is configured on the platform, outside the questionnaire schema."
        )
        for pool in pools:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(f"{pool['pool_name']}: ").bold = True
            p.add_run(
                f"shows {pool['show_count']} of {len(pool['pool_pages'])} pages "
                f"({', '.join(pool['pool_pages'])})."
            )
            if pool.get("condition_var") and pool["show_count"] == 1:
                cmap = pool.get("condition_map") or {}
                mapping = ", ".join(
                    f"{page} → “{cmap.get(page, page)}”" for page in pool["pool_pages"]
                )
                p.add_run(f" Sets variable “{pool['condition_var']}”: {mapping}.")

    for page in pages:
        doc.add_page_break()
        h = doc.add_heading(loc_text(page.get("title")) or page.get("name", ""), level=1)
        for run in h.runs:
            run.font.color.rgb = NAVY
        _note(doc, f"Internal page name: {page.get('name', '?')}")
        if page.get("visibleIf"):
            _note(doc, "Shown only if: " + humanize_expr(page["visibleIf"]))
        pool = _pool_for_page(pools, page.get("name", ""))
        if pool:
            _note(doc, _pool_page_note(pool, page["name"]))
        if page.get("description"):
            p = doc.add_paragraph()
            run = p.add_run(loc_text(page["description"]))
            run.italic = True
        for el in page.get("elements", []):
            _render_element(doc, el, survey_locales)

    import io
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def pools_from_rows(rows) -> list:
    """Adapt rand_pools DB rows to the dicts build_review_docx expects."""
    pools = []
    for r in rows:
        cmap = None
        if r["condition_map"]:
            try:
                cmap = json.loads(r["condition_map"])
            except (json.JSONDecodeError, TypeError):
                cmap = None
        pools.append({
            "pool_name": r["pool_name"],
            "pool_pages": json.loads(r["pool_pages"]),
            "show_count": r["show_count"],
            "condition_var": r["condition_var"],
            "condition_map": cmap,
        })
    return pools
